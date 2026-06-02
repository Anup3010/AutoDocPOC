# """
# AutoDoc Generator
# Generates project documentation and PR impact analysis using HuggingFace Inference API.
# Triggered automatically on git commit.
# """

# import os
# import json
# import time
# import hashlib
# from datetime import datetime
# from typing import Dict, Optional, Tuple
# from dataclasses import dataclass, asdict

# from doc_generator.git_utils import build_diff_result, get_full_project_code, GitDiffResult
# from doc_generator.code_parser import parse_project, format_module_summary

# # Load .env only when running locally
# # On GitHub Actions secrets come from environment directly
# # if os.path.exists(".env"):
# #     try:
# #         from dotenv import load_dotenv
# #         load_dotenv()
# #         print("📁 Local mode — loaded from .env file")
# #     except ImportError:
# #         # Manual load if python-dotenv not installed
# #         with open(".env") as f:
# #             for line in f:
# #                 line = line.strip()
# #                 if line and not line.startswith("#") and "=" in line:
# #                     key, _, val = line.partition("=")
# #                     os.environ.setdefault(key.strip(), val.strip())
# # else:
# #     print("☁️  CI mode — using GitHub Actions secrets")

# # ─────────────────────────────────────────────────────
# # Configuration
# # ─────────────────────────────────────────────────────
# HF_API_KEY = os.environ.get("HF_API_KEY", "")
# HF_MODEL = "mistralai/Mistral-7B-Instruct-v0.3"   # Free tier model
# GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "")
# GEMINI_MODEL = "gemini-2.5-flash"   # free, fast, good quality
# # Alternatives: "HuggingFaceH4/zephyr-7b-beta", "microsoft/Phi-3-mini-4k-instruct"

# DOCS_OUTPUT_DIR = os.environ.get("AUTODOC_OUTPUT_DIR", "generated_docs")
# MAX_CODE_CHARS = 8000   # Keep prompts under token limits
# MAX_DIFF_CHARS = 4000


# @dataclass
# class TokenUsage:
#     prompt_tokens: int
#     completion_tokens: int
#     total_tokens: int
#     estimated_cost_usd: float   # Rough HF estimate


# @dataclass
# class GenerationResult:
#     success: bool
#     full_doc_path: Optional[str]
#     pr_doc_path: Optional[str]
#     token_usage: Optional[TokenUsage]
#     error: Optional[str]
#     generation_time_sec: float


# # ─────────────────────────────────────────────────────
# # Token estimation utilities
# # ─────────────────────────────────────────────────────


    
    
# def estimate_tokens(text: str) -> int:
#     """
#     Rough token estimation: ~4 chars per token for English/code.
#     HuggingFace free tier doesn't always return exact counts.
#     """
#     return max(1, len(text) // 4)


# def estimate_cost(input_tokens: int, output_tokens: int, model: str = HF_MODEL) -> float:
#     """
#     Rough cost estimate. HF Inference API free tier = $0.
#     For reference: Mistral-7B on HF Pro ≈ $0.0004/1K tokens.
#     """
#     cost_per_1k = 0.0004
#     return (input_tokens + output_tokens) / 1000 * cost_per_1k


# # ─────────────────────────────────────────────────────
# # HuggingFace API Call
# # ─────────────────────────────────────────────────────
# def call_huggingface_api(prompt: str, max_tokens: int = 1500) -> Tuple[str, TokenUsage]:
#     """
#     Call Google Gemini API to generate documentation.
#     Drop-in replacement for HuggingFace — same inputs and outputs.
#     """
#     import requests

#     if not GEMINI_API_KEY:
#         raise ValueError(
#             "GEMINI_API_KEY not set. "
#             "Get your free key from: https://aistudio.google.com"
#         )

#     url = (
#         f"https://generativelanguage.googleapis.com/v1beta/models/"
#         f"{GEMINI_MODEL}:generateContent?key={GEMINI_API_KEY}"
#     )
#     headers = {"Content-Type": "application/json"}
#     payload = {
#         "contents": [{"parts": [{"text": prompt}]}],
#         "generationConfig": {
#             "maxOutputTokens": max_tokens,
#             "temperature": 0.3
#         }
#     }

#     response = requests.post(url, headers=headers, json=payload, timeout=60)

#     if response.status_code != 200:
#         raise RuntimeError(
#             f"Gemini API error {response.status_code}: {response.text[:300]}"
#         )

#     result = response.json()

#     # Extract generated text
#     text = result["candidates"][0]["content"]["parts"][0]["text"]

#     # Extract token usage (Gemini returns exact counts)
#     usage_meta = result.get("usageMetadata", {})
#     prompt_tokens = usage_meta.get("promptTokenCount", estimate_tokens(prompt))
#     output_tokens = usage_meta.get("candidatesTokenCount", estimate_tokens(text))
#     total_tokens = usage_meta.get("totalTokenCount", prompt_tokens + output_tokens)

#     usage = TokenUsage(
#         prompt_tokens=prompt_tokens,
#         completion_tokens=output_tokens,
#         total_tokens=total_tokens,
#         estimated_cost_usd=0.0    # Gemini free tier = $0
#     )
#     return text, usage

# # def call_huggingface_api(prompt: str, max_tokens: int = 1500) -> Tuple[str, TokenUsage]:
#     """
#     Call HuggingFace Inference API using requests (no SDK needed).
    
#     Returns:
#         (generated_text, TokenUsage)
#     """
#     import requests

#     if not HF_API_KEY:
#         raise ValueError(
#             "HF_API_KEY environment variable not set. "
#             "Get your key from: https://huggingface.co/settings/tokens"
#         )

#     url = f"https://api-inference.huggingface.co/models/{HF_MODEL}"
#     headers = {
#         "Authorization": f"Bearer {HF_API_KEY}",
#         "Content-Type": "application/json"
#     }

#     # Format for instruction models
#     formatted_prompt = f"<s>[INST] {prompt} [/INST]"

#     payload = {
#         "inputs": formatted_prompt,
#         "parameters": {
#             "max_new_tokens": max_tokens,
#             "temperature": 0.3,
#             "do_sample": True,
#             "return_full_text": False,
#             "stop": ["</s>", "[INST]"]
#         },
#         "options": {
#             "wait_for_model": True,
#             "use_cache": False
#         }
#     }

#     input_tokens = estimate_tokens(formatted_prompt)

#     response = requests.post(url, headers=headers, json=payload, timeout=120)

#     if response.status_code == 503:
#         # Model is loading, wait and retry
#         time.sleep(20)
#         response = requests.post(url, headers=headers, json=payload, timeout=120)

#     if response.status_code != 200:
#         raise RuntimeError(
#             f"HuggingFace API error {response.status_code}: {response.text[:500]}"
#         )

#     result = response.json()

#     if isinstance(result, list) and len(result) > 0:
#         generated_text = result[0].get("generated_text", "")
#     elif isinstance(result, dict):
#         generated_text = result.get("generated_text", "")
#     else:
#         generated_text = str(result)

#     # Strip instruction echo if present
#     if "[/INST]" in generated_text:
#         generated_text = generated_text.split("[/INST]", 1)[-1].strip()

#     output_tokens = estimate_tokens(generated_text)
#     usage = TokenUsage(
#         prompt_tokens=input_tokens,
#         completion_tokens=output_tokens,
#         total_tokens=input_tokens + output_tokens,
#         estimated_cost_usd=estimate_cost(input_tokens, output_tokens)
#     )

#     return generated_text, usage


# # ─────────────────────────────────────────────────────
# # Prompt Builders
# # ─────────────────────────────────────────────────────
# def build_full_doc_prompt(project_summary: str, project_name: str) -> str:
#     """Build prompt for full project documentation."""
#     return f"""You are a senior software architect writing comprehensive technical documentation.

# Generate a complete, professional project documentation for the following Python project.
# Make it human-readable, well-structured with clear sections, and useful for a new developer joining the team.

# PROJECT NAME: {project_name}

# PROJECT CODE STRUCTURE:
# {project_summary[:MAX_CODE_CHARS]}

# Write the documentation with these sections:
# 1. **Project Overview** - What this project does and its purpose
# 2. **Architecture Overview** - High-level structure and design patterns used
# 3. **API Endpoints** - All REST endpoints with parameters, request/response examples
# 4. **Data Models** - Description of all models and their fields
# 5. **Service Layer** - Business logic in each service with key methods explained
# 6. **Business Rules** - Important validation rules and constraints
# 7. **Error Handling** - How errors are handled across layers
# 8. **Getting Started** - How a new developer can run and test this project
# 9. **Dependencies & Key Libraries** - What external libraries are used and why

# Be specific, mention actual class names, method names, and endpoint paths.
# Write in clear English that both technical and semi-technical readers can understand."""


# def build_pr_doc_prompt(
#     diff_result: GitDiffResult,
#     project_summary: str,
#     project_name: str
# ) -> str:
#     """Build prompt for PR/commit impact analysis document."""
#     commit = diff_result.commit
#     diff_details = []
#     for f in diff_result.changed_files[:5]:
#         status_map = {"A": "ADDED", "M": "MODIFIED", "D": "DELETED"}
#         status = status_map.get(f.status, f.status)
#         diff_details.append(
#             f"File: {f.path} [{status}] +{f.additions}/-{f.deletions} lines\n"
#             f"Diff snippet:\n{f.diff[:MAX_DIFF_CHARS // max(len(diff_result.changed_files), 1)]}"
#         )

#     diff_text = "\n\n".join(diff_details) if diff_details else "No Python files were modified."

#     return f"""You are a senior software engineer reviewing a pull request for a Python project.

# Analyze the following code changes and write a professional PR Review Document that will help a senior developer understand the impact and decide whether to approve or reject the pull request.

# PROJECT: {project_name}
# COMMIT: {commit.short_hash} by {commit.author} on {commit.date}
# BRANCH: {commit.branch}
# COMMIT MESSAGE: "{commit.message}"
# CHANGES SUMMARY: {diff_result.summary}
# Total: +{diff_result.total_additions} lines added, -{diff_result.total_deletions} lines deleted

# PROJECT CONTEXT (existing code):
# {project_summary[:MAX_CODE_CHARS // 2]}

# CODE CHANGES:
# {diff_text}

# Write a PR Review Document with these sections:
# 1. **Change Summary** - What changed in plain English (2-3 sentences max)
# 2. **Files Changed** - Table of files modified with what changed in each
# 3. **Impact Analysis** - What other parts of the codebase are affected by these changes
# 4. **API Impact** - Are any API endpoints added, changed, or broken?
# 5. **Breaking Changes** - List any breaking changes (if none, say "None detected")
# 6. **Business Logic Changes** - Any changes to validation rules, business logic
# 7. **Risk Assessment** - Low/Medium/High risk with justification
# 8. **Testing Recommendations** - What test cases should be written or run
# 9. **Reviewer Checklist** - 5-8 specific items the senior developer should verify
# 10. **Recommendation** - Approve / Request Changes / Reject with reason

# Be specific about actual class names, methods, and endpoints. Be direct and concise."""


# # ─────────────────────────────────────────────────────
# # Document Savers
# # ─────────────────────────────────────────────────────
# def save_as_markdown(content: str, filepath: str):
#     """Save document as Markdown file."""
#     os.makedirs(os.path.dirname(filepath), exist_ok=True)
#     with open(filepath, "w", encoding="utf-8") as f:
#         f.write(content)


# def save_as_docx(content: str, filepath: str, title: str):
#     """Save document as Word .docx file."""
#     try:
#         from docx import Document
#         from docx.shared import Pt, RGBColor
#         from docx.enum.text import WD_ALIGN_PARAGRAPH
#         import re

#         doc = Document()

#         # Title
#         title_para = doc.add_heading(title, 0)
#         title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

#         # Process content
#         lines = content.split("\n")
#         for line in lines:
#             line = line.rstrip()
#             if line.startswith("## "):
#                 doc.add_heading(line[3:], level=2)
#             elif line.startswith("### "):
#                 doc.add_heading(line[4:], level=3)
#             elif line.startswith("# "):
#                 doc.add_heading(line[2:], level=1)
#             elif line.startswith("**") and line.endswith("**") and len(line) > 4:
#                 p = doc.add_paragraph()
#                 run = p.add_run(line.strip("*"))
#                 run.bold = True
#             elif line.startswith("- ") or line.startswith("* "):
#                 doc.add_paragraph(line[2:], style="List Bullet")
#             elif re.match(r'^\d+\. ', line):
#                 doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style="List Number")
#             elif line.startswith("|") and "|" in line[1:]:
#                 # Table row (simple handling)
#                 doc.add_paragraph(line)
#             elif line.strip() == "" or line.strip() == "---":
#                 continue
#             else:
#                 p = doc.add_paragraph()
#                 # Handle bold within paragraph
#                 parts = re.split(r'(\*\*[^*]+\*\*)', line)
#                 for part in parts:
#                     if part.startswith("**") and part.endswith("**"):
#                         run = p.add_run(part[2:-2])
#                         run.bold = True
#                     else:
#                         p.add_run(part)

#         doc.save(filepath)
#     except ImportError:
#         # Fallback: save as text with .docx extension note
#         with open(filepath.replace(".docx", ".txt"), "w") as f:
#             f.write(f"=== {title} ===\n\n{content}")


# def save_as_pdf(content: str, filepath: str, title: str):
#     """Save document as PDF using reportlab."""
#     try:
#         from reportlab.lib.pagesizes import A4
#         from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
#         from reportlab.lib.units import cm
#         from reportlab.lib import colors
#         from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
#         from reportlab.lib.enums import TA_LEFT, TA_CENTER
#         import re

#         doc = SimpleDocTemplate(
#             filepath,
#             pagesize=A4,
#             rightMargin=2*cm, leftMargin=2*cm,
#             topMargin=2*cm, bottomMargin=2*cm
#         )
#         styles = getSampleStyleSheet()

#         # Custom styles
#         title_style = ParagraphStyle(
#             'CustomTitle',
#             parent=styles['Title'],
#             fontSize=18,
#             textColor=colors.HexColor('#1a365d'),
#             spaceAfter=12,
#             alignment=TA_CENTER
#         )
#         h1_style = ParagraphStyle(
#             'CustomH1', parent=styles['Heading1'],
#             fontSize=14, textColor=colors.HexColor('#2c5282'),
#             spaceBefore=12, spaceAfter=6
#         )
#         h2_style = ParagraphStyle(
#             'CustomH2', parent=styles['Heading2'],
#             fontSize=12, textColor=colors.HexColor('#2b6cb0'),
#             spaceBefore=8, spaceAfter=4
#         )
#         body_style = ParagraphStyle(
#             'CustomBody', parent=styles['Normal'],
#             fontSize=9, leading=14, spaceAfter=4
#         )
#         bullet_style = ParagraphStyle(
#             'CustomBullet', parent=styles['Normal'],
#             fontSize=9, leading=14, leftIndent=20, spaceAfter=2,
#             bulletIndent=10
#         )

#         story = []
#         story.append(Paragraph(title, title_style))
#         story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor('#2c5282')))
#         story.append(Spacer(1, 12))

#         for line in content.split("\n"):
#             line = line.rstrip()
#             # Escape HTML special chars for reportlab
#             safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

#             if not safe.strip():
#                 story.append(Spacer(1, 4))
#             elif safe.startswith("## ") or safe.startswith("**") and safe.endswith("**"):
#                 text = re.sub(r'^#+\s*', '', safe).strip("*")
#                 story.append(Paragraph(text, h1_style))
#             elif safe.startswith("### "):
#                 story.append(Paragraph(safe[4:], h2_style))
#             elif safe.startswith("# "):
#                 story.append(Paragraph(safe[2:], h1_style))
#             elif safe.startswith("- ") or safe.startswith("* "):
#                 story.append(Paragraph(f"• {safe[2:]}", bullet_style))
#             elif re.match(r'^\d+\.', safe):
#                 story.append(Paragraph(safe, bullet_style))
#             elif safe.strip() == "---":
#                 story.append(HRFlowable(width="100%", thickness=0.5))
#             else:
#                 # Handle inline bold
#                 safe = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', safe)
#                 try:
#                     story.append(Paragraph(safe, body_style))
#                 except Exception:
#                     story.append(Paragraph(line[:200], body_style))

#         doc.build(story)
#     except ImportError:
#         # Fallback
#         with open(filepath.replace(".pdf", ".txt"), "w") as f:
#             f.write(f"=== {title} ===\n\n{content}")


# # ─────────────────────────────────────────────────────
# # Main Generator
# # ─────────────────────────────────────────────────────
# def push_to_confluence(markdown_content: str, commit_info: dict):
#     """
#     Push generated documentation to Confluence page automatically.
#     Requires: CONFLUENCE_URL, CONFLUENCE_USER, CONFLUENCE_TOKEN, 
#               CONFLUENCE_PAGE_ID in .env
#     """
#     import requests
#     import base64

#     # Read Confluence config from .env
#     confluence_url  = os.environ.get("CONFLUENCE_URL", "")
#     confluence_user = os.environ.get("CONFLUENCE_USER", "")
#     confluence_token = os.environ.get("CONFLUENCE_TOKEN", "")
#     page_id         = os.environ.get("CONFLUENCE_PAGE_ID", "")

#     if not all([confluence_url, confluence_user, confluence_token, page_id]):
#         print("⚠️  Confluence config missing — skipping Confluence update")
#         return

#     # Convert markdown to Confluence storage format (basic HTML)
#     import markdown2
#     html_content = markdown2.markdown(markdown_content)

#     # Get current page version (Confluence requires version number)
#     auth = (confluence_user, confluence_token)
#     page_url = f"{confluence_url}/rest/api/content/{page_id}"
#     current = requests.get(page_url, auth=auth).json()
#     current_version = current["version"]["number"]

#     # Update the page
#     payload = {
#         "version": {"number": current_version + 1},
#         "title": f"Project Documentation — {commit_info['short_hash']}",
#         "type": "page",
#         "body": {
#             "storage": {
#                 "value": html_content,
#                 "representation": "storage"
#             }
#         }
#     }

#     response = requests.put(
#         page_url,
#         json=payload,
#         auth=auth,
#         headers={"Content-Type": "application/json"}
#     )

#     if response.status_code == 200:
#         print(f"✅ Confluence page updated: {confluence_url}/wiki/pages/{page_id}")
#     else:
#         print(f"⚠️  Confluence update failed: {response.status_code}")


# def generate_documents(repo_path: str = ".") -> GenerationResult:
#     """
#     Main entry point: generates both documents after a git commit.
    
#     Returns:
#         GenerationResult with paths to generated documents and token usage
#     """
#     start_time = time.time()
#     os.makedirs(DOCS_OUTPUT_DIR, exist_ok=True)

#     # 1. Gather git info
#     print("📂 Analyzing git repository...")
#     diff_result = build_diff_result(repo_path)
#     if not diff_result:
#         return GenerationResult(
#             success=False, full_doc_path=None, pr_doc_path=None,
#             token_usage=None, error="Could not read git repository",
#             generation_time_sec=time.time() - start_time
#         )

#     commit = diff_result.commit
#     print(f"✅ Commit: {commit.short_hash} — {commit.message}")

#     # 2. Parse project code
#     print("🔍 Parsing project code...")
#     project_code = get_full_project_code(repo_path)
#     parsed = parse_project(project_code)
    
#     project_name = os.path.basename(os.path.abspath(repo_path))
#     project_summary = "\n\n".join(
#         format_module_summary(m) for m in list(parsed.values())[:15]
#     )

#     total_token_usage = TokenUsage(0, 0, 0, 0.0)

#     # ─── Document 1: Full Project Documentation ───
#     print("📝 Generating full project documentation...")
#     full_doc_prompt = build_full_doc_prompt(project_summary, project_name)

#     try:
#         full_doc_content, usage1 = call_huggingface_api(full_doc_prompt, max_tokens=2000)
#         total_token_usage.prompt_tokens += usage1.prompt_tokens
#         total_token_usage.completion_tokens += usage1.completion_tokens
#         total_token_usage.total_tokens += usage1.total_tokens
#         total_token_usage.estimated_cost_usd += usage1.estimated_cost_usd
#         print(f"   Tokens: {usage1.prompt_tokens} in / {usage1.completion_tokens} out")
#     except Exception as e:
#         return GenerationResult(
#             success=False, full_doc_path=None, pr_doc_path=None,
#             token_usage=None, error=f"Full doc generation failed: {str(e)}",
#             generation_time_sec=time.time() - start_time
#         )

#     # Add header
#     full_doc_header = (
#         f"# {project_name} — Project Documentation\n\n"
#         f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  \n"
#         f"**Latest Commit:** `{commit.short_hash}` — {commit.message}  \n"
#         f"**Author:** {commit.author} | **Branch:** {commit.branch}  \n"
#         f"**Token Usage:** {usage1.prompt_tokens} input / {usage1.completion_tokens} output "
#         f"(~{usage1.total_tokens} total)\n\n---\n\n"
#     )
#     full_doc_text = full_doc_header + full_doc_content

#     # Save full doc
#     timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
#     full_doc_md = os.path.join(DOCS_OUTPUT_DIR, f"project_doc_{timestamp}.md")
#     full_doc_docx = os.path.join(DOCS_OUTPUT_DIR, f"project_doc_{timestamp}.docx")
#     full_doc_pdf = os.path.join(DOCS_OUTPUT_DIR, f"project_doc_{timestamp}.pdf")

#     save_as_markdown(full_doc_text, full_doc_md)
#     save_as_docx(full_doc_text, full_doc_docx, f"{project_name} — Project Documentation")
#     save_as_pdf(full_doc_text, full_doc_pdf, f"{project_name} — Project Documentation")
#     print(f"✅ Full doc saved: {full_doc_md}")

#     # ─── Document 2: PR / Commit Impact Analysis ───
#     print("🔍 Generating PR impact analysis...")
#     pr_doc_prompt = build_pr_doc_prompt(diff_result, project_summary, project_name)

#     try:
#         pr_doc_content, usage2 = call_huggingface_api(pr_doc_prompt, max_tokens=1500)
#         total_token_usage.prompt_tokens += usage2.prompt_tokens
#         total_token_usage.completion_tokens += usage2.completion_tokens
#         total_token_usage.total_tokens += usage2.total_tokens
#         total_token_usage.estimated_cost_usd += usage2.estimated_cost_usd
#         print(f"   Tokens: {usage2.prompt_tokens} in / {usage2.completion_tokens} out")
#     except Exception as e:
#         pr_doc_content = f"⚠️ PR analysis generation failed: {str(e)}"
#         usage2 = TokenUsage(0, 0, 0, 0.0)

#     pr_doc_header = (
#         f"# PR Impact Analysis — `{commit.short_hash}`\n\n"
#         f"**Commit:** {commit.short_hash} | **Branch:** {commit.branch}  \n"
#         f"**Author:** {commit.author} ({commit.email})  \n"
#         f"**Date:** {commit.date}  \n"
#         f"**Message:** *{commit.message}*  \n"
#         f"**Changes:** +{diff_result.total_additions} / -{diff_result.total_deletions} lines  \n"
#         f"**Token Usage:** {usage2.prompt_tokens} input / {usage2.completion_tokens} output "
#         f"(~{usage2.total_tokens} total)\n\n---\n\n"
#     )
#     pr_doc_text = pr_doc_header + pr_doc_content

#     pr_doc_md = os.path.join(DOCS_OUTPUT_DIR, f"pr_analysis_{commit.short_hash}_{timestamp}.md")
#     pr_doc_docx = os.path.join(DOCS_OUTPUT_DIR, f"pr_analysis_{commit.short_hash}_{timestamp}.docx")
#     pr_doc_pdf = os.path.join(DOCS_OUTPUT_DIR, f"pr_analysis_{commit.short_hash}_{timestamp}.pdf")

#     save_as_markdown(pr_doc_text, pr_doc_md)
#     save_as_docx(pr_doc_text, pr_doc_docx, f"PR Impact Analysis — {commit.short_hash}")
#     save_as_pdf(pr_doc_text, pr_doc_pdf, f"PR Impact Analysis — {commit.short_hash}")
#     print(f"✅ PR doc saved: {pr_doc_md}")

#     # Save metadata JSON for the viewer
#     meta = {
#         "timestamp": timestamp,
#         "commit": {
#             "hash": commit.hash,
#             "short_hash": commit.short_hash,
#             "message": commit.message,
#             "author": commit.author,
#             "date": commit.date,
#             "branch": commit.branch
#         },
#         "token_usage": {
#             "prompt_tokens": total_token_usage.prompt_tokens,
#             "completion_tokens": total_token_usage.completion_tokens,
#             "total_tokens": total_token_usage.total_tokens,
#             "estimated_cost_usd": round(total_token_usage.estimated_cost_usd, 6)
#         },
#         "files": {
#             "full_doc_md": full_doc_md,
#             "full_doc_docx": full_doc_docx,
#             "full_doc_pdf": full_doc_pdf,
#             "pr_doc_md": pr_doc_md,
#             "pr_doc_docx": pr_doc_docx,
#             "pr_doc_pdf": pr_doc_pdf,
#         },
#         "changes_summary": diff_result.summary,
#         "files_changed": len(diff_result.changed_files)
#     }
#     meta_path = os.path.join(DOCS_OUTPUT_DIR, f"meta_{timestamp}.json")
#     with open(meta_path, "w") as f:
#         json.dump(meta, f, indent=2)

#     elapsed = time.time() - start_time
#     print(f"\n🎉 Done! Generated in {elapsed:.1f}s")
#     print(f"📊 Total tokens: {total_token_usage.prompt_tokens} in / "
#           f"{total_token_usage.completion_tokens} out = {total_token_usage.total_tokens} total")
#     print(f"💰 Estimated cost: ${total_token_usage.estimated_cost_usd:.6f}")

#     return GenerationResult(
#         success=True,
#         full_doc_path=full_doc_md,
#         pr_doc_path=pr_doc_md,
#         token_usage=total_token_usage,
#         error=None,
#         generation_time_sec=elapsed
#     )


# if __name__ == "__main__":
#     import sys
#     repo = sys.argv[1] if len(sys.argv) > 1 else "."
#     result = generate_documents(repo)
#     if not result.success:
#         print(f"❌ Error: {result.error}")
#         sys.exit(1)


"""
AutoDoc Generator
Generates project documentation and PR impact analysis using Gemini AI.
Triggered automatically on git commit via post-commit hook.
"""

import os
import json
import time
from datetime import datetime
from typing import Dict, Optional, Tuple
from dataclasses import dataclass

from doc_generator.git_utils import build_diff_result, get_full_project_code, GitDiffResult
from doc_generator.code_parser import parse_project, format_module_summary

# ─────────────────────────────────────────────────────
# Load .env manually — works locally and on GitHub Actions
# ─────────────────────────────────────────────────────
def _load_env():
    env_path = os.path.join(os.getcwd(), ".env")
    if os.path.exists(env_path):
        with open(env_path, encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith("#") and "=" in line:
                    key, _, val = line.partition("=")
                    os.environ.setdefault(key.strip(), val.strip())
        print("📁 Local mode — loaded from .env file")
    else:
        print("☁️  CI mode — using environment variables")

_load_env()

# ─────────────────────────────────────────────────────
# Configuration
# ─────────────────────────────────────────────────────
GEMINI_API_KEY     = os.environ.get("GEMINI_API_KEY", "")
GEMINI_MODEL       = "gemini-2.5-flash"
CONFLUENCE_URL     = os.environ.get("CONFLUENCE_URL", "")
CONFLUENCE_USER    = os.environ.get("CONFLUENCE_USER", "")
CONFLUENCE_TOKEN   = os.environ.get("CONFLUENCE_TOKEN", "")
CONFLUENCE_PAGE_ID = os.environ.get("CONFLUENCE_PAGE_ID", "")
AUTODOC_MODE       = os.environ.get("AUTODOC_MODE", "both")

DOCS_OUTPUT_DIR = os.environ.get("AUTODOC_OUTPUT_DIR", "generated_docs")
MAX_CODE_CHARS  = 6000
MAX_DIFF_CHARS  = 3000


# ─────────────────────────────────────────────────────
# Data Classes
# ─────────────────────────────────────────────────────
@dataclass
class TokenUsage:
    prompt_tokens: int
    completion_tokens: int
    total_tokens: int
    estimated_cost_usd: float


@dataclass
class GenerationResult:
    success: bool
    full_doc_path: Optional[str]
    pr_doc_path: Optional[str]
    token_usage: Optional[TokenUsage]
    error: Optional[str]
    generation_time_sec: float


# ─────────────────────────────────────────────────────
# Token Utilities
# ─────────────────────────────────────────────────────
def estimate_tokens(text: str) -> int:
    return max(1, len(text) // 4)


# ─────────────────────────────────────────────────────
# Gemini API Call
# ─────────────────────────────────────────────────────
def call_huggingface_api(prompt: str, max_tokens: int = 8192) -> Tuple[str, TokenUsage]:
    """Call Google Gemini API to generate documentation."""
    import requests

    if not GEMINI_API_KEY:
        raise ValueError(
            "GEMINI_API_KEY not set. "
            "Get your free key from: https://aistudio.google.com"
        )

    url = (
        f"https://generativelanguage.googleapis.com/v1beta/models/"
        f"{GEMINI_MODEL}:generateContent?key={GEMINI_API_KEY}"
    )
    headers = {"Content-Type": "application/json"}
    payload = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {
            "maxOutputTokens": max_tokens,
            "temperature": 0.3,
            "topP": 0.9
        }
    }

    response = requests.post(url, headers=headers, json=payload, timeout=120)

    if response.status_code != 200:
        raise RuntimeError(
            f"Gemini API error {response.status_code}: {response.text[:300]}"
        )

    result = response.json()
    text = result["candidates"][0]["content"]["parts"][0]["text"]

    usage_meta = result.get("usageMetadata", {})
    prompt_tokens  = usage_meta.get("promptTokenCount", estimate_tokens(prompt))
    output_tokens  = usage_meta.get("candidatesTokenCount", estimate_tokens(text))
    total_tokens   = usage_meta.get("totalTokenCount", prompt_tokens + output_tokens)

    usage = TokenUsage(
        prompt_tokens=prompt_tokens,
        completion_tokens=output_tokens,
        total_tokens=total_tokens,
        estimated_cost_usd=0.0
    )
    return text, usage


# ─────────────────────────────────────────────────────
# Prompt Builders
# ─────────────────────────────────────────────────────
def build_full_doc_prompt(project_summary: str, project_name: str) -> str:
    return f"""You are a senior software architect writing comprehensive technical documentation.

Generate a complete, professional project documentation for the following Python project.
Make it human-readable, well-structured with clear sections, and useful for a new developer joining the team.

PROJECT NAME: {project_name}

PROJECT CODE STRUCTURE:
{project_summary[:MAX_CODE_CHARS]}

Write the documentation with these sections:
1. **Project Overview** - What this project does and its purpose
2. **Architecture Overview** - High-level structure and design patterns used
3. **API Endpoints** - All REST endpoints with parameters and examples
4. **Data Models** - Description of all models and their fields
5. **Service Layer** - Business logic in each service with key methods explained
6. **Business Rules** - Important validation rules and constraints
7. **Error Handling** - How errors are handled across layers
8. **Getting Started** - How a new developer can run and test this project
9. **Dependencies & Key Libraries** - What external libraries are used and why

Be specific, mention actual class names, method names, and endpoint paths.
Write at least 800 words. Be detailed and thorough.
Write in clear English that both technical and semi-technical readers can understand."""


def build_pr_doc_prompt(
    diff_result: GitDiffResult,
    project_summary: str,
    project_name: str
) -> str:
    commit = diff_result.commit
    diff_details = []
    for f in diff_result.changed_files[:5]:
        status_map = {"A": "ADDED", "M": "MODIFIED", "D": "DELETED"}
        status = status_map.get(f.status, f.status)
        diff_details.append(
            f"File: {f.path} [{status}] +{f.additions}/-{f.deletions} lines\n"
            f"Diff:\n{f.diff[:MAX_DIFF_CHARS // max(len(diff_result.changed_files), 1)]}"
        )

    diff_text = "\n\n".join(diff_details) if diff_details else "No source files modified."

    return f"""You are a senior software engineer reviewing a pull request.

Analyze the following code changes and write a professional PR Review Document.

PROJECT: {project_name}
COMMIT: {commit.short_hash} by {commit.author} on {commit.date}
BRANCH: {commit.branch}
COMMIT MESSAGE: "{commit.message}"
CHANGES SUMMARY: {diff_result.summary}
Total: +{diff_result.total_additions} lines added, -{diff_result.total_deletions} lines deleted

PROJECT CONTEXT:
{project_summary[:MAX_CODE_CHARS // 2]}

CODE CHANGES:
{diff_text}

Write a detailed PR Review Document with these sections:
1. **Change Summary** - What changed in plain English
2. **Files Changed** - List of files with what changed in each
3. **Impact Analysis** - What other parts of the codebase are affected
4. **API Impact** - Are any API endpoints added, changed, or broken?
5. **Breaking Changes** - List any breaking changes (if none, say "None detected")
6. **Business Logic Changes** - Any changes to validation rules or business logic
7. **Risk Assessment** - Low/Medium/High risk with justification
8. **Testing Recommendations** - What test cases should be written or run
9. **Reviewer Checklist** - 5-8 specific items the senior developer should verify
10. **Recommendation** - Approve / Request Changes / Reject with reason

Be specific about actual class names, methods, and endpoints.
Write at least 500 words. Be detailed."""


# ─────────────────────────────────────────────────────
# Document Savers
# ─────────────────────────────────────────────────────
def save_as_markdown(content: str, filepath: str):
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)


def save_as_docx(content: str, filepath: str, title: str):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re

        doc = Document()
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for line in content.split("\n"):
            line = line.rstrip()
            if line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif re.match(r'^\d+\. ', line):
                doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style="List Number")
            elif line.strip() in ("", "---"):
                continue
            else:
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*[^*]+\*\*)', line)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        p.add_run(part[2:-2]).bold = True
                    else:
                        p.add_run(part)
        doc.save(filepath)
    except ImportError:
        with open(filepath.replace(".docx", ".txt"), "w", encoding="utf-8") as f:
            f.write(f"=== {title} ===\n\n{content}")


def save_as_pdf(content: str, filepath: str, title: str):
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.enums import TA_CENTER
        import re

        doc = SimpleDocTemplate(
            filepath, pagesize=A4,
            rightMargin=2*cm, leftMargin=2*cm,
            topMargin=2*cm, bottomMargin=2*cm
        )
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            'T', parent=styles['Title'],
            fontSize=18, textColor=colors.HexColor('#1a365d'),
            spaceAfter=12, alignment=TA_CENTER
        )
        h1_style = ParagraphStyle(
            'H1', parent=styles['Heading1'],
            fontSize=14, textColor=colors.HexColor('#2c5282'),
            spaceBefore=12, spaceAfter=6
        )
        h2_style = ParagraphStyle(
            'H2', parent=styles['Heading2'],
            fontSize=12, textColor=colors.HexColor('#2b6cb0'),
            spaceBefore=8, spaceAfter=4
        )
        body_style = ParagraphStyle(
            'B', parent=styles['Normal'],
            fontSize=9, leading=14, spaceAfter=4
        )
        bullet_style = ParagraphStyle(
            'BL', parent=styles['Normal'],
            fontSize=9, leading=14,
            leftIndent=20, spaceAfter=2, bulletIndent=10
        )

        story = [
            Paragraph(title, title_style),
            HRFlowable(width="100%", thickness=2, color=colors.HexColor('#2c5282')),
            Spacer(1, 12)
        ]

        for line in content.split("\n"):
            line = line.rstrip()
            safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            if not safe.strip():
                story.append(Spacer(1, 4))
            elif safe.startswith("## "):
                story.append(Paragraph(safe[3:], h1_style))
            elif safe.startswith("### "):
                story.append(Paragraph(safe[4:], h2_style))
            elif safe.startswith("# "):
                story.append(Paragraph(safe[2:], h1_style))
            elif safe.startswith("- ") or safe.startswith("* "):
                story.append(Paragraph(f"• {safe[2:]}", bullet_style))
            elif re.match(r'^\d+\.', safe):
                story.append(Paragraph(safe, bullet_style))
            elif safe.strip() == "---":
                story.append(HRFlowable(width="100%", thickness=0.5))
            else:
                safe = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', safe)
                try:
                    story.append(Paragraph(safe, body_style))
                except Exception:
                    story.append(Paragraph(line[:200], body_style))

        doc.build(story)
    except ImportError:
        with open(filepath.replace(".pdf", ".txt"), "w", encoding="utf-8") as f:
            f.write(f"=== {title} ===\n\n{content}")


# ─────────────────────────────────────────────────────
# Confluence Integration
# ─────────────────────────────────────────────────────
def push_to_confluence(markdown_content: str, commit_info: dict):
    """Push full project documentation to Confluence page."""
    import requests
    import markdown2

    if not all([CONFLUENCE_URL, CONFLUENCE_USER, CONFLUENCE_TOKEN, CONFLUENCE_PAGE_ID]):
        print("⚠️  Confluence credentials missing in .env — skipping Confluence update")
        return

    print("📤 Pushing to Confluence...")

    try:
        html_content = markdown2.markdown(
            markdown_content,
            extras=["tables", "fenced-code-blocks"]
        )

        auth = (CONFLUENCE_USER, CONFLUENCE_TOKEN)
        page_url = f"{CONFLUENCE_URL}/rest/api/content/{CONFLUENCE_PAGE_ID}"

        current = requests.get(page_url, auth=auth, timeout=30).json()
        current_version = current["version"]["number"]
        page_title = current["title"]

        payload = {
            "version": {"number": current_version + 1},
            "title": page_title,
            "type": "page",
            "body": {
                "storage": {
                    "value": html_content,
                    "representation": "storage"
                }
            }
        }

        response = requests.put(
            page_url,
            json=payload,
            auth=auth,
            headers={"Content-Type": "application/json"},
            timeout=30
        )

        if response.status_code == 200:
            print(f"✅ Confluence updated successfully")
        else:
            print(f"⚠️  Confluence update failed: {response.status_code} — {response.text[:200]}")

    except Exception as e:
        print(f"⚠️  Confluence error: {e}")


# ─────────────────────────────────────────────────────
# Main Generator
# ─────────────────────────────────────────────────────
def generate_documents(repo_path: str = ".") -> GenerationResult:
    """
    Main entry point.
    Reads AUTODOC_MODE from environment:
      pr_only  = feature branch → only PR analysis
      both     = main/master    → full doc + PR analysis + Confluence update
    """
    start_time = time.time()
    os.makedirs(DOCS_OUTPUT_DIR, exist_ok=True)

    # Re-read mode in case it was set after module load
    mode = os.environ.get("AUTODOC_MODE", AUTODOC_MODE)
    print(f"📋 AutoDoc Mode: {mode}")

    # ── Gather git info ──────────────────────────────
    print("📂 Analyzing git repository...")
    diff_result = build_diff_result(repo_path)
    if not diff_result:
        return GenerationResult(
            success=False, full_doc_path=None, pr_doc_path=None,
            token_usage=None, error="Could not read git repository",
            generation_time_sec=time.time() - start_time
        )

    commit = diff_result.commit
    print(f"✅ Commit: {commit.short_hash} — {commit.message}")
    print(f"🌿 Branch: {commit.branch}")

    # ── Parse project code ───────────────────────────
    print("🔍 Parsing project code...")
    project_code = get_full_project_code(repo_path)
    parsed = parse_project(project_code)
    project_name = os.path.basename(os.path.abspath(repo_path))
    project_summary = "\n\n".join(
        format_module_summary(m) for m in list(parsed.values())[:15]
    )

    total_token_usage = TokenUsage(0, 0, 0, 0.0)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    full_doc_md = None
    full_doc_docx = None
    full_doc_pdf = None
    pr_doc_md = None
    pr_doc_docx = None
    pr_doc_pdf = None

    # ── Document 1: Full Project Doc (main/master only) ──
    if mode in ("both", "full_only"):
        print("📝 Generating full project documentation...")
        full_doc_prompt = build_full_doc_prompt(project_summary, project_name)

        try:
            full_doc_content, usage1 = call_huggingface_api(
                full_doc_prompt, max_tokens=8192
            )
            total_token_usage.prompt_tokens    += usage1.prompt_tokens
            total_token_usage.completion_tokens += usage1.completion_tokens
            total_token_usage.total_tokens      += usage1.total_tokens
            print(f"   Tokens: {usage1.prompt_tokens} in / {usage1.completion_tokens} out")

            full_doc_header = (
                f"# {project_name} — Project Documentation\n\n"
                f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  \n"
                f"**Latest Commit:** `{commit.short_hash}` — {commit.message}  \n"
                f"**Author:** {commit.author} | **Branch:** {commit.branch}  \n"
                f"**Tokens:** {usage1.prompt_tokens} in / {usage1.completion_tokens} out\n\n---\n\n"
            )
            full_doc_text = full_doc_header + full_doc_content

            full_doc_md   = os.path.join(DOCS_OUTPUT_DIR, f"project_doc_{timestamp}.md")
            full_doc_docx = os.path.join(DOCS_OUTPUT_DIR, f"project_doc_{timestamp}.docx")
            full_doc_pdf  = os.path.join(DOCS_OUTPUT_DIR, f"project_doc_{timestamp}.pdf")

            save_as_markdown(full_doc_text, full_doc_md)
            save_as_docx(full_doc_text, full_doc_docx, f"{project_name} — Project Documentation")
            save_as_pdf(full_doc_text, full_doc_pdf, f"{project_name} — Project Documentation")
            print(f"✅ Full doc saved: {full_doc_md}")

            # Push to Confluence after full doc is saved
            push_to_confluence(full_doc_text, {
                "short_hash": commit.short_hash,
                "message": commit.message
            })

        except Exception as e:
            print(f"❌ Full doc generation failed: {e}")

    else:
        print("⏭️  Skipping full doc — feature branch (only generates on main/master)")

    # ── Document 2: PR Analysis (every commit) ───────
    print("🔍 Generating PR impact analysis...")
    pr_doc_prompt = build_pr_doc_prompt(diff_result, project_summary, project_name)

    try:
        pr_doc_content, usage2 = call_huggingface_api(
            pr_doc_prompt, max_tokens=8192
        )
        total_token_usage.prompt_tokens    += usage2.prompt_tokens
        total_token_usage.completion_tokens += usage2.completion_tokens
        total_token_usage.total_tokens      += usage2.total_tokens
        print(f"   Tokens: {usage2.prompt_tokens} in / {usage2.completion_tokens} out")

    except Exception as e:
        pr_doc_content = f"⚠️ PR analysis generation failed: {e}"
        usage2 = TokenUsage(0, 0, 0, 0.0)

    pr_doc_header = (
        f"# PR Impact Analysis — `{commit.short_hash}`\n\n"
        f"**Commit:** {commit.short_hash} | **Branch:** {commit.branch}  \n"
        f"**Author:** {commit.author} ({commit.email})  \n"
        f"**Date:** {commit.date}  \n"
        f"**Message:** *{commit.message}*  \n"
        f"**Changes:** +{diff_result.total_additions} / -{diff_result.total_deletions} lines  \n"
        f"**Tokens:** {usage2.prompt_tokens} in / {usage2.completion_tokens} out\n\n---\n\n"
    )
    pr_doc_text = pr_doc_header + pr_doc_content

    pr_doc_md   = os.path.join(DOCS_OUTPUT_DIR, f"pr_analysis_{commit.short_hash}_{timestamp}.md")
    pr_doc_docx = os.path.join(DOCS_OUTPUT_DIR, f"pr_analysis_{commit.short_hash}_{timestamp}.docx")
    pr_doc_pdf  = os.path.join(DOCS_OUTPUT_DIR, f"pr_analysis_{commit.short_hash}_{timestamp}.pdf")

    save_as_markdown(pr_doc_text, pr_doc_md)
    save_as_docx(pr_doc_text, pr_doc_docx, f"PR Impact Analysis — {commit.short_hash}")
    save_as_pdf(pr_doc_text, pr_doc_pdf, f"PR Impact Analysis — {commit.short_hash}")
    print(f"✅ PR doc saved: {pr_doc_md}")

    # ── Save metadata JSON ───────────────────────────
    meta = {
        "timestamp": timestamp,
        "mode": mode,
        "commit": {
            "hash": commit.hash,
            "short_hash": commit.short_hash,
            "message": commit.message,
            "author": commit.author,
            "date": commit.date,
            "branch": commit.branch
        },
        "token_usage": {
            "prompt_tokens": total_token_usage.prompt_tokens,
            "completion_tokens": total_token_usage.completion_tokens,
            "total_tokens": total_token_usage.total_tokens,
            "estimated_cost_usd": 0.0
        },
        "files": {
            "full_doc_md":   full_doc_md,
            "full_doc_docx": full_doc_docx,
            "full_doc_pdf":  full_doc_pdf,
            "pr_doc_md":     pr_doc_md,
            "pr_doc_docx":   pr_doc_docx,
            "pr_doc_pdf":    pr_doc_pdf,
        },
        "changes_summary": diff_result.summary,
        "files_changed":   len(diff_result.changed_files)
    }

    meta_path = os.path.join(DOCS_OUTPUT_DIR, f"meta_{timestamp}.json")
    with open(meta_path, "w", encoding="utf-8") as f:
        json.dump(meta, f, indent=2)

    elapsed = time.time() - start_time
    print(f"\n🎉 Done! Generated in {elapsed:.1f}s")
    print(f"📊 Tokens: {total_token_usage.prompt_tokens} in / "
          f"{total_token_usage.completion_tokens} out = {total_token_usage.total_tokens} total")

    return GenerationResult(
        success=True,
        full_doc_path=full_doc_md,
        pr_doc_path=pr_doc_md,
        token_usage=total_token_usage,
        error=None,
        generation_time_sec=elapsed
    )


if __name__ == "__main__":
    import sys
    repo = sys.argv[1] if len(sys.argv) > 1 else "."
    result = generate_documents(repo)
    if not result.success:
        print(f"❌ Error: {result.error}")
        sys.exit(1)